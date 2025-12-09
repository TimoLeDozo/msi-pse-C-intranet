import os
import json
import ollama
from datetime import datetime
from fastapi import FastAPI, Request, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from typing import Optional, Dict, Any

# --- CONFIGURATION ---
# Modèle puissant pour le serveur (assure-toi de l'avoir pull : ollama pull mixtral)
MODEL_NAME = "mixtral"  
TEMPLATE_PATH = "model/Template_B2.docx"
OUTPUT_DIR = "output"

app = FastAPI()

# Montage des dossiers statiques
app.mount("/static", StaticFiles(directory="static"), name="static")
app.mount("/assets", StaticFiles(directory="assets"), name="assets")
templates = Jinja2Templates(directory="templates")

# Création du dossier de sortie si absent
os.makedirs(OUTPUT_DIR, exist_ok=True)

# --- MODÈLE DE DONNÉES (Ce que le JS envoie) ---
class PropaleData(BaseModel):
    # Adapter ces champs selon les "name" exacts de ton formulaire HTML
    titre: str = "Projet Sans Titre"
    entrepriseNom: str = "Client Inconnu"
    clientNom: str = ""
    contexte: str = ""      # Champ 'ia_probleme' du front
    solution: str = ""      # Champ 'ia_solution' du front
    objectifs: str = ""     # Champ 'ia_objectifs' du front
    # Accepte tout autre champ dynamique
    extra: Dict[str, Any] = {}

    class Config:
        extra = "allow" # Permet de recevoir tout le JSON du front sans planter

# --- LOGIQUE MÉTIER ---

def generate_text_with_ollama(data: dict):
    """Génère le contenu rédactionnel via Ollama"""
    
    system_prompt = """
    Tu es un expert en ingénierie d'affaires Icam. 
    Ton but est de rédiger les sections d'une proposition commerciale au format JSON.
    Sois professionnel, convaincant et structuré.
    """
    
    user_prompt = f"""
    CLIENT: {data.get('entrepriseNom')}
    CONTEXTE: {data.get('ia_probleme')}
    SOLUTION PROPOSÉE: {data.get('ia_solution')}
    OBJECTIFS: {data.get('ia_objectifs')}
    
    Rédige 3 sections :
    1. "contexte_detaille" : Analyse de la situation actuelle.
    2. "demarche_methodo" : Méthodologie proposée (Phasage, Agile, etc.).
    3. "phrase_accroche" : Une phrase percutante pour l'intro.

    Réponds UNIQUEMENT un JSON valide sous la forme :
    {{
        "contexte_detaille": "...",
        "demarche_methodo": "...",
        "phrase_accroche": "..."
    }}
    """

    print(f"🤖 Appel Ollama ({MODEL_NAME})...")
    try:
        response = ollama.chat(model=MODEL_NAME, messages=[
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': user_prompt},
        ], format='json')
        
        return json.loads(response['message']['content'])
    except Exception as e:
        print(f"❌ Erreur Ollama: {e}")
        return {
            "contexte_detaille": "Erreur de génération IA.",
            "demarche_methodo": "Erreur de génération IA.",
            "phrase_accroche": "Erreur."
        }

def fill_docx(data_source: dict, ai_content: dict):
    """Remplit le template Word"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    safe_name = "".join([c for c in data_source.get('entrepriseNom', 'Client') if c.isalnum()])
    filename = f"Propale_{safe_name}_{timestamp}.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    
    doc = Document(TEMPLATE_PATH)
    
    # Fusion des données manuelles et IA
    full_data = {**data_source, **ai_content}
    
    # Remplacement simple (Placeholder type {{cle}})
    for paragraph in doc.paragraphs:
        for key, value in full_data.items():
            if isinstance(value, str) and f"{{{{{key}}}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", value)
                
    # Idem pour les tableaux (simplifié)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in full_data.items():
                        if isinstance(value, str) and f"{{{{{key}}}}}" in p.text:
                            p.text = p.text.replace(f"{{{{{key}}}}}", value)

    doc.save(output_path)
    return filename

# --- ROUTES API ---

@app.get("/")
async def serve_interface(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/api/generate")
async def api_generate(request: Request):
    data = await request.json()
    
    # 1. Génération IA
    ai_content = generate_text_with_ollama(data)
    
    # 2. Remplissage Word
    filename = fill_docx(data, ai_content)
    
    # 3. Réponse au format attendu par le JS de l'Hypothèse B
    return {
        "success": True,
        "url": f"/download/{filename}", # Lien de téléchargement
        "cost": {"total": 0, "currency": "EUR"},
        "ai_debug": ai_content
    }

@app.get("/download/{filename}")
async def download_file(filename: str):
    file_path = os.path.join(OUTPUT_DIR, filename)
    if os.path.exists(file_path):
        return FileResponse(file_path, filename=filename)
    raise HTTPException(status_code=404, detail="Fichier introuvable")

if __name__ == "__main__":
    import uvicorn
    # host="0.0.0.0" rend le serveur accessible sur le réseau intranet
    uvicorn.run(app, host="0.0.0.0", port=8000)